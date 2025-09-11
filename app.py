import io
import os
import shutil
import zipfile
import tempfile
import pandas as pd
import streamlit as st
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.utils.cell import coordinate_from_string, column_index_from_string

from PIL import Image

# ==============================
# Configuración inicial de la página
# ==============================
st.set_page_config(
    page_title="Generador DOCUMENTACIÓN CLIENTES - EssilorLuxottica",
    layout="wide"
)

# ==============================
# CSS personalizado
# ==============================
st.markdown(
    """
    <style>
    .stApp {
        background-color: #000000; /* Fondo negro */
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ==============================
# Mostrar imágenes lado a lado
# ==============================
col1, col2, col3 = st.columns([2, 1, 1])

with col1:
    logo = Image.open("logo.png")
    st.image(logo, width=500)

with col3:
    otra_imagen = Image.open("servioptica.png")
    st.image(otra_imagen, width=300)

# ==============================
# Títulos
# ==============================
st.title("📂 Generador DOCUMENTACIÓN CLIENTES — Autor : KEVIN EDUARDO GUTIERREZ CASTILLO")

st.markdown(
    """
    **Instrucciones rápidas**:
    1. Sube **DATOS CLIENTES** (Excel donde se encuentran los datos de todos los clientes).
    2. Sube **FR-EI-02** (ENTREGA DE EQUIPO A CONFORMIDAD Y CONDICIONES DE GARANTÍA - Word).
    3. Sube **FR-EI-04**, **FR-EI-03**, **FR-EI-05** (Excel: HOJA DE VIDA, PROTOCOLO DE MANTENIMIENTO, CRONOGRAMA).
    4. Haz clic en **Generar** — se descargará `DOCUMENTACION_CLIENTES.zip` con un ZIP por cliente/equipo.
    """
)
st.markdown("---")

# -------------------------
# Upload inputs
# -------------------------
st.header("1) Cargar archivos base")
col1, col2 = st.columns([2, 1])

with col1:
    plantilla_datos_file = st.file_uploader("📊 DATOS CLIENTES", type=["xlsx", "xls", "csv"])
    fr_ei_02_file = st.file_uploader("📝 FR-EI-02 ENTREGA DE EQUIPO (Word)", type=["docx"])
    fr_ei_04_file = st.file_uploader("📑 FR-EI-04 HOJA DE VIDA DEL EQUIPO", type=["xlsx"])
    fr_ei_03_file = st.file_uploader("📑 FR-EI-03 PROTOCOLO DE MANTENIMIENTO", type=["xlsx"])
    cronograma_file = st.file_uploader("📑 FR-EI-05 CRONOGRAMA DE MANTENIMIENTO", type=["xlsx"])

with col2:
    st.markdown("**Estado de carga**")
    st.write("DATOS CLIENTES:", getattr(plantilla_datos_file, "name", "—"))
    st.write("FR-EI-02 (Word):", getattr(fr_ei_02_file, "name", "—"))
    st.write("FR-EI-04 (Excel):", getattr(fr_ei_04_file, "name", "—"))
    st.write("FR-EI-03 (Excel):", getattr(fr_ei_03_file, "name", "—"))
    st.write("FR-EI-05 (Excel):", getattr(cronograma_file, "name", "—"))

st.markdown("---")

# -------------------------
# Helpers
# -------------------------
def write_respecting_merged(ws, target_cell, value):
    """
    Función general para escribir en celdas mergeadas o normales.
    """
    for merged in ws.merged_cells.ranges:
        if target_cell in str(merged):
            tl_col = get_column_letter(merged.min_col)
            tl_row = merged.min_row
            ws[f"{tl_col}{tl_row}"] = value
            return
    ws[target_cell] = value

# Función mejorada SOLO para F10 y G10
def write_respecting_merged_f10_g10(ws, target_cell, value):
    col, row = coordinate_from_string(target_cell)
    col_idx = column_index_from_string(col)

    for merged in ws.merged_cells.ranges:
        if (merged.min_col <= col_idx <= merged.max_col) and (merged.min_row <= row <= merged.max_row):
            ws.cell(merged.min_row, merged.min_col, value)
            return
    ws[target_cell] = value

def safe_str(val):
    return "" if pd.isna(val) else str(val)

# -------------------------
# Validación mínima
# -------------------------
if not (plantilla_datos_file and fr_ei_02_file and fr_ei_04_file and fr_ei_03_file and cronograma_file):
    st.info("Sube todos los archivos listados arriba para habilitar la generación.")
    st.stop()

# -------------------------
# Leer plantilla_datos en DataFrame
# -------------------------
try:
    if plantilla_datos_file.name.lower().endswith(".csv"):
        df = pd.read_csv(plantilla_datos_file)
    else:
        df = pd.read_excel(plantilla_datos_file, engine="openpyxl")
except Exception as e:
    st.error(f"No se pudo leer PLANTILLA_DATOS: {e}")
    st.stop()

st.subheader("Vista previa (primeras filas)")
st.dataframe(df.head(10))

# Leer bytes de plantillas
fr_ei_02_bytes = fr_ei_02_file.read()
fr_ei_04_bytes = fr_ei_04_file.read()
fr_ei_03_bytes = fr_ei_03_file.read()
cronograma_bytes = cronograma_file.read()

# -------------------------
# Botón de generación
# -------------------------
st.header("2) Generar documentación")
base_name = st.text_input("Nombre base para archivos/ZIP (opcional)", value="DOCUMENTACION_CLIENTES")
generate = st.button("🚀 Generar DOCUMENTACIÓN")

if generate:
    with st.spinner("Generando archivos..."):
        tmp_root = tempfile.mkdtemp()
        output_dir = os.path.join(tmp_root, "DOCUMENTACION_CLIENTES")
        os.makedirs(output_dir, exist_ok=True)
        created_zip_paths = []

        # Iterar sobre filas de clientes/equipos
        for index, row in df.iterrows():
            try:
                cliente = safe_str(row.get("CLIENTE", "")).strip().replace(" ", "_")
                equipo = safe_str(row.get("NOMBRE DEL EQUIPO", "")).strip().replace(" ", "_")
                referencia = safe_str(row.get("REFERENCIA", ""))
                serie = safe_str(row.get("SERIE", ""))
                fecha_word = safe_str(row.get("FECHA INSTALACION(WORD)", ""))
                nit_cliente = safe_str(row.get("NIT CLIENTE", ""))
                tipo_mmto = safe_str(row.get("TIPO DE MANTENIMIENTO", ""))
                frecuencia = safe_str(row.get("FRECUENCIA", ""))
                direccion = safe_str(row.get("DIRECCION", ""))
                modelo = safe_str(row.get("MODELO", ""))
                ubicacion_equipo = safe_str(row.get("UBICACIÓN DEL EQUIPO (ÁREA)", ""))
                dia = safe_str(row.get("DD", ""))
                mes = safe_str(row.get("MM", ""))
                anio = safe_str(row.get("AA", ""))
                entidad = safe_str(row.get("ENTIDAD", ""))
                ciudad = safe_str(row.get("CIUDAD", ""))
                telefono = safe_str(row.get("TELEFONO CLIENTE", ""))

                # Crear carpeta temporal por cliente-equipo
                folder_name = f"{cliente}_{equipo}" if cliente or equipo else f"fila_{index}"
                folder_path = os.path.join(output_dir, folder_name)
                os.makedirs(folder_path, exist_ok=True)

                # ---------- A) FR-EI-02 Word ----------
                doc = Document(io.BytesIO(fr_ei_02_bytes))
                for p in doc.paragraphs:
                    if "(Nombre cliente)" in p.text:
                        p.text = p.text.replace("(Nombre cliente)", safe_str(row.get("CLIENTE", "")))
                try:
                    table = doc.tables[0]
                    table.cell(0, 1).text = safe_str(row.get("NOMBRE DEL EQUIPO", ""))
                    table.cell(1, 1).text = referencia
                    table.cell(2, 1).text = serie
                    table.cell(3, 1).text = fecha_word
                except Exception as e:
                    st.warning(f"No se pudo llenar tabla FR-EI-02 fila {index}: {e}")
                word_name = f"FR-EI-02-{cliente}-{equipo}.docx"
                doc.save(os.path.join(folder_path, word_name))

                # ---------- B) FR-EI-04 Hoja de vida ----------
                wb_hv = load_workbook(io.BytesIO(fr_ei_04_bytes))
                ws_hv = wb_hv.active
                try:
                    ws_hv["D9"] = row.get("NOMBRE DEL EQUIPO", "")
                    ws_hv["V24"] = dia
                    ws_hv["X24"] = mes
                    ws_hv["Y24"] = anio
                    ws_hv["D22"] = entidad
                    ws_hv["AE22"] = ciudad
                    ws_hv["AE24"] = telefono
                    ws_hv["AD7"] = serie
                except Exception as e:
                    st.warning(f"Error FR-EI-04 fila {index}: {e}")
                wb_hv.save(os.path.join(folder_path, f"FR-EI-04-{cliente}-{equipo}.xlsx"))

                # ---------- C) FR-EI-03 Protocolo MMTO ----------
                wb_mmto = load_workbook(io.BytesIO(fr_ei_03_bytes))
                ws_mmto = wb_mmto.active
                try:
                    ws_mmto["A12"] = row.get("NOMBRE DEL EQUIPO", "")
                except Exception as e:
                    st.warning(f"Error FR-EI-03 fila {index}: {e}")
                wb_mmto.save(os.path.join(folder_path, f"FR-EI-03-{cliente}-{equipo}.xlsx"))

                # ---------- D) FR-EI-05 Cronograma ----------
                wb_crono = load_workbook(io.BytesIO(cronograma_bytes))
                ws_crono = wb_crono.active
                try:
                    write_respecting_merged(ws_crono, "B10", row.get("NOMBRE DEL EQUIPO", ""))
                    write_respecting_merged(ws_crono, "E10", serie)

                    # 👇 SOLO ESTAS DOS USAN LA NUEVA LÓGICA
                    write_respecting_merged_f10_g10(ws_crono, "F10", tipo_mmto)   # Tipo de mantenimiento
                    write_respecting_merged_f10_g10(ws_crono, "G10", frecuencia)  # Frecuencia de mantenimiento

                    write_respecting_merged(ws_crono, "D5", nit_cliente)
                    write_respecting_merged(ws_crono, "R6", fecha_word)
                    write_respecting_merged(ws_crono, "F5", direccion)
                    write_respecting_merged(ws_crono, "R5", ciudad)
                    write_respecting_merged(ws_crono, "D10", modelo)
                    write_respecting_merged(ws_crono, "D6", ubicacion_equipo)
                    write_respecting_merged(ws_crono, "D4", row.get("CLIENTE", ""))
                    write_respecting_merged(ws_crono, "R4", telefono)
                except Exception as e:
                    st.warning(f"Error FR-EI-05 fila {index}: {e}")
                wb_crono.save(os.path.join(folder_path, f"FR-EI-05-{cliente}-{equipo}.xlsx"))

                # ---------- ZIP por cliente_equipo ----------
                zip_path = os.path.join(output_dir, f"{cliente}_{equipo}.zip")
                with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zipf:
                    for fname in os.listdir(folder_path):
                        zipf.write(os.path.join(folder_path, fname), arcname=fname)
                shutil.rmtree(folder_path)
                created_zip_paths.append(zip_path)

            except Exception as e:
                st.error(f"Error procesando fila {index}: {e}")
                continue

        # ---------- ZIP final ----------
        final_zip_path = os.path.join(tmp_root, f"{base_name}.zip")
        with zipfile.ZipFile(final_zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for p in created_zip_paths:
                zf.write(p, arcname=os.path.basename(p))

        with open(final_zip_path, "rb") as f:
            data_bytes = f.read()

        st.success("✅ Proceso terminado. Descarga el ZIP con todos los clientes.")
        st.download_button("⬇️ Descargar DOCUMENTACION_CLIENTES.zip", data=data_bytes, file_name=f"{base_name}.zip", mime="application/zip")
